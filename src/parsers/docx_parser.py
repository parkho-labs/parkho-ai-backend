import os
from typing import Dict, Any
from docx import Document

from .base_parser import BaseContentParser, ContentParseResult


class DOCXParser(BaseContentParser):
    def __init__(self):
        self.max_file_size = 5 * 1024 * 1024  # 5MB limit

    async def parse(self, source: str, **kwargs) -> ContentParseResult:
        try:
            if not os.path.exists(source):
                return ContentParseResult("", error=f"DOCX file not found: {source}")

            file_size = os.path.getsize(source)
            if file_size > self.max_file_size:
                return ContentParseResult(
                    "",
                    error=f"DOCX file too large: {file_size} bytes (max: {self.max_file_size})"
                )

            document = Document(source)
            content_parts = []
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)

            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            content_parts.append(text)

            content = "\n\n".join(content_parts)

            if not content.strip():
                return ContentParseResult("", error="No text content found in DOCX")

            metadata = self._extract_metadata(document, source)

            return ContentParseResult(
                content=content.strip(),
                title=metadata.get("title"),
                metadata=metadata
            )

        except Exception as e:
            return ContentParseResult("", error=f"Failed to parse DOCX: {str(e)}")

    def _extract_metadata(self, document: Document, file_path: str) -> Dict[str, Any]:
        metadata = {
            "file_name": os.path.basename(file_path),
            "file_size": os.path.getsize(file_path),
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
        }

        try:
            core_props = document.core_properties
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.created:
                metadata["created"] = core_props.created.isoformat()
            if core_props.modified:
                metadata["modified"] = core_props.modified.isoformat()
        except Exception:
            pass

        if "title" not in metadata:
            metadata["title"] = os.path.splitext(os.path.basename(file_path))[0]

        return metadata

    def supports_source(self, source: str) -> bool:
        return source.lower().endswith((".docx", ".doc"))

    @property
    def supported_types(self) -> list[str]:
        return ["docx", "doc"]